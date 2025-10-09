#!/usr/bin/env python3
"""
AI e-Book Generator (Self-Citing Edition) — Robust
--------------------------------------------------
Plans -> drafts -> assembles a full book from a single idea.

Highlights
- Self-cite mode: per-chapter Crossref search, strict in-text citations [@key],
  and an auto-built References chapter (no hallucinated citations).
- JSON planning is enforced with OpenAI Responses API + json_schema.
- Hardened for GPT-5 quirks (max_tokens vs max_completion_tokens; temperature).
- UTC timestamp fix, retries, and guards against empty plan/chapters.

Quick start (OpenAI):
    python -m pip install 'pydantic==2.*' 'rich==13.*' 'tqdm==4.*' 'tenacity==8.*' 'python-dateutil==2.*' 'requests==2.*'
    python -m pip install 'openai==1.*'
    # optional outputs:
    python -m pip install 'ebooklib==0.*' 'markdown==3.*' 'python-docx==1.*' 'beautifulsoup4==4.*' 'reportlab==4.*'
    # optional: if supplying --bibtex
    python -m pip install 'bibtexparser==1.*'

Run:
    export OPENAI_API_KEY="sk-..."
    python ebook_generator.py \
      --idea "A microbiologist uses AI to decode the hidden ecology of dairy farm microbes" \
      --genre "textbook" \
      --audience "upper-undergrad and graduate students" \
      --reading-level "university" \
      --style "scholarly, rigorous, precise" \
      --min-words 20000 --max-words 30000 \
      --backend openai --model gpt-5-thinking \
      --auto-cite \
      --outdir ./book_out
"""
from __future__ import annotations

import argparse
import json
import os
import re
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from pydantic import BaseModel, Field
from rich.console import Console
from tqdm import tqdm

# ----- Optional imports guarded at runtime -----
try:
    import openai  # type: ignore
except Exception:
    openai = None

try:
    import anthropic  # type: ignore
except Exception:
    anthropic = None

try:
    import requests  # HTTP (Ollama + Crossref queries)
except Exception:
    requests = None

try:
    from ebooklib import epub  # type: ignore
except Exception:
    epub = None

try:
    import markdown  # type: ignore
except Exception:
    markdown = None

try:
    import docx  # type: ignore
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    import bibtexparser  # type: ignore
except Exception:
    bibtexparser = None

console = Console()

# =========================
# Helpers (JSON parsing/repair)
# =========================
def _extract_json_block(text: str) -> str:
    """Extract the first JSON object from text, allowing fenced code blocks."""
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.S | re.I)
    if m:
        return m.group(1)
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]
    raise ValueError("No JSON object found")

def _parse_or_repair_json(raw: str) -> dict:
    """Attempt to parse or minimally repair model output to strict JSON."""
    # 1) direct
    try:
        return json.loads(raw)
    except Exception:
        pass
    # 2) extract block
    try:
        return json.loads(_extract_json_block(raw))
    except Exception:
        pass
    # 3) strip trailing commas ", }" or ", ]"
    cleaned = re.sub(r",\s*([}\]])", r"\1", raw)
    try:
        return json.loads(_extract_json_block(cleaned))
    except Exception:
        pass
    raise ValueError("Could not parse JSON")

# =========================
# LLM Backends
# =========================
class LLMError(Exception):
    pass

class LLMClient:
    def complete(self, system: str, prompt: str, *, temperature: float = 1.0, max_tokens: int = 1500, model: str = "") -> str:
        raise NotImplementedError

class OpenAIClient(LLMClient):
    """Tolerant client that adapts to GPT-5-style params; also supports structured JSON via Responses API."""
    def __init__(self, model: str):
        if openai is None:
            raise RuntimeError("openai package not installed. Run: pip install openai==1.*")
        self.model = model
        self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    @retry(wait=wait_exponential(min=1, max=60), stop=stop_after_attempt(6), reraise=True,
           retry=retry_if_exception_type(Exception))
    def complete(self, system: str, prompt: str, *, temperature: float = 1.0,
                 max_tokens: int = 1500, model: str = "") -> str:
        m = model or self.model
        base = {
            "model": m,
            "messages": [{"role": "system", "content": system},
                         {"role": "user", "content": prompt}],
        }

        def chat_call(extra):
            return self.client.chat.completions.create(**{**base, **extra})

        # First try legacy params (works for many models)
        try:
            resp = chat_call({"max_tokens": max_tokens, "temperature": temperature})
            return resp.choices[0].message.content or ""
        except Exception as e:
            msg = str(e)

            # Retry if model requires max_completion_tokens
            if "max_completion_tokens" in msg:
                try:
                    resp = chat_call({"extra_body": {"max_completion_tokens": max_tokens},
                                      "temperature": temperature})
                    return resp.choices[0].message.content or ""
                except Exception as e2:
                    msg = str(e2)

            # Retry if model disallows temperature (only default supported)
            if "temperature" in msg.lower() and "unsupported" in msg.lower():
                try:
                    # omit temperature entirely
                    resp = chat_call({"max_tokens": max_tokens})
                    return resp.choices[0].message.content or ""
                except Exception as e3:
                    msg = str(e3)

            # Final fallback: Responses API (max_output_tokens) and omit temp
            try:
                resp2 = self.client.responses.create(
                    model=m,
                    input=[
                        {"role": "system", "content": system},
                        {"role": "user", "content": prompt},
                    ],
                    max_output_tokens=max_tokens,
                )
                return getattr(resp2, "output_text", None) or ""
            except Exception:
                pass

            raise LLMError(msg)

    def structured(self, system: str, prompt: str, *, json_schema: dict, max_tokens: int = 2000, model: str = "") -> dict:
        """
        Ask the model to return JSON that MUST validate against json_schema.
        Uses the Responses API with response_format=json_schema.
        """
        m = model or self.model
        try:
            resp = self.client.responses.create(
                model=m,
                input=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                response_format={
                    "type": "json_schema",
                    "json_schema": {
                        "name": json_schema.get("$id", "schema"),
                        "schema": json_schema,
                        "strict": True,
                    },
                },
                max_output_tokens=max_tokens,
            )
            txt = getattr(resp, "output_text", None) or ""
            return json.loads(txt) if txt else {}
        except Exception as e:
            raise LLMError(f"structured() failed: {e}")

class AnthropicClient(LLMClient):
    def __init__(self, model: str):
        if anthropic is None:
            raise RuntimeError("anthropic package not installed. Run: pip install anthropic==0.*")
        self.model = model
        self.client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    @retry(wait=wait_exponential(min=1, max=60), stop=stop_after_attempt(6), reraise=True,
           retry=retry_if_exception_type(Exception))
    def complete(self, system: str, prompt: str, *, temperature: float = 1.0, max_tokens: int = 1500, model: str = "") -> str:
        m = model or self.model
        try:
            msg = self.client.messages.create(
                model=m,
                max_tokens=max_tokens,
                temperature=temperature,
                system=system,
                messages=[{"role": "user", "content": prompt}],
            )
            parts = []
            for block in msg.content:
                if getattr(block, "type", "") == "text":
                    parts.append(block.text)
            return "\n".join(parts)
        except Exception as e:
            raise LLMError(str(e))

class OllamaClient(LLMClient):
    def __init__(self, model: str, host: str = "http://localhost:11434"):
        if requests is None:
            raise RuntimeError("requests not installed. Run: pip install requests==2.*")
        self.model = model
        self.host = host.rstrip("/")

    @retry(wait=wait_exponential(min=1, max=30), stop=stop_after_attempt(6), reraise=True,
           retry=retry_if_exception_type(Exception))
    def complete(self, system: str, prompt: str, *, temperature: float = 1.0, max_tokens: int = 1500, model: str = "") -> str:
        m = model or self.model
        try:
            payload = {
                "model": m,
                "options": {"temperature": temperature, "num_predict": max_tokens},
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                "stream": False,
            }
            r = requests.post(f"{self.host}/api/chat", json=payload, timeout=180)
            r.raise_for_status()
            data = r.json()
            return data.get("message", {}).get("content", "")
        except Exception as e:
            raise LLMError(str(e))

# =========================
# Data Models
# =========================
class BookMeta(BaseModel):
    title: str
    subtitle: str
    author: str
    genre: str
    audience: str
    reading_level: str
    style: str
    promise: str = Field(description="Unique value proposition / hook.")

class SectionPlan(BaseModel):
    title: str
    summary: str

class ChapterPlan(BaseModel):
    number: int
    title: str
    purpose: str
    key_points: List[str]
    target_words: int
    sections: List[SectionPlan]

class BookPlan(BaseModel):
    synopsis: str
    parts: List[str]
    chapters: List[ChapterPlan]
    total_target_words: int

# =========================
# Prompts & Schemas
# =========================
SYSTEM_POLICY = (
    "You are a seasoned book author and editor. Write original, coherent, factual prose. "
    "Avoid harmful or hateful content. Do not fabricate real citations. Maintain consistency."
)

META_PROMPT = """
Given the book idea below, propose professional metadata and a strong market hook.
Return ONLY valid JSON with keys: title, subtitle, author, genre, audience, reading_level, style, promise.

Book idea: {idea}
Preferred genre: {genre}
Target audience: {audience}
Reading level: {reading_level}
Style notes: {style}
""".strip()

PLAN_PROMPT = """
Design a complete long-form book based on the metadata and idea below.
Return ONLY valid JSON with keys:
- synopsis (str),
- parts (list[str]),
- chapters (list of objects with: number(int), title(str), purpose(str), key_points(list[str]),
  target_words(int), sections(list of {{title, summary}})),
- total_target_words (int).

Constraints:
- 16–22 chapters.
- Each chapter 2,500–3,000 words (we will enforce during drafting).
- No prose outside the JSON.

Metadata: {meta_json}
Book idea: {idea}
Target total words: {min_words}–{max_words}
""".strip()

META_SCHEMA = {
    "$schema": "https://json-schema.org/draft/2020-12/schema",
    "$id": "meta.schema.json",
    "type": "object",
    "required": ["title", "subtitle", "author", "genre", "audience", "reading_level", "style", "promise"],
    "properties": {
        "title": {"type": "string", "minLength": 1},
        "subtitle": {"type": "string"},
        "author": {"type": "string"},
        "genre": {"type": "string"},
        "audience": {"type": "string"},
        "reading_level": {"type": "string"},
        "style": {"type": "string"},
        "promise": {"type": "string"}
    },
    "additionalProperties": False
}

PLAN_SCHEMA = {
    "$schema": "https://json-schema.org/draft/2020-12/schema",
    "$id": "plan.schema.json",
    "type": "object",
    "required": ["synopsis", "parts", "chapters", "total_target_words"],
    "properties": {
        "synopsis": {"type": "string"},
        "parts": {"type": "array", "items": {"type": "string"}},
        "total_target_words": {"type": "integer", "minimum": 10000},
        "chapters": {
            "type": "array",
            "minItems": 8,
            "items": {
                "type": "object",
                "required": ["number", "title", "purpose", "key_points", "target_words", "sections"],
                "properties": {
                    "number": {"type": "integer", "minimum": 1},
                    "title": {"type": "string", "minLength": 1},
                    "purpose": {"type": "string"},
                    "key_points": {"type": "array", "items": {"type": "string"}, "minItems": 3},
                    "target_words": {"type": "integer", "minimum": 1200},
                    "sections": {
                        "type": "array",
                        "minItems": 3,
                        "items": {
                            "type": "object",
                            "required": ["title", "summary"],
                            "properties": {
                                "title": {"type": "string"},
                                "summary": {"type": "string"}
                            },
                            "additionalProperties": False
                        }
                    }
                },
                "additionalProperties": False
            }
        }
    },
    "additionalProperties": False
}

# =========================
# Auto-citation (Crossref)
# =========================
CROSSREF_API = "https://api.crossref.org/works"

@dataclass
class Ref:
    key: str
    title: str
    authors: str
    year: str
    venue: str
    doi: str = ""
    url: str = ""

    def to_bibtex(self) -> str:
        fields = {"title": self.title, "author": self.authors, "year": self.year}
        if self.venue: fields["journal"] = self.venue
        if self.doi: fields["doi"] = self.doi
        if self.url: fields["url"] = self.url
        esc = lambda x: x.replace("&", "\\&")
        body = ",\n  ".join(f"{k} = {{{esc(v)}}}" for k, v in fields.items() if v)
        return f"@article{{{self.key},\n  {body}\n}}\n"

def _http_get(url: str, params: Dict[str, Any]) -> Dict[str, Any]:
    if requests is None:
        raise RuntimeError("requests not installed (needed for auto-cite).")
    headers = {"User-Agent": os.getenv("AUTOCITE_USER_AGENT", "AutoCite/1.0 (mailto:example@example.com)")}
    r = requests.get(url, params=params, headers=headers, timeout=30)
    r.raise_for_status()
    return r.json()

def search_crossref(query: str, rows: int = 5, min_year: int = 2000) -> List[Ref]:
    data = _http_get(CROSSREF_API, {
        "query": query,
        "rows": rows,
        "sort": "is-referenced-by-count",
        "order": "desc",
        "filter": f"type:journal-article,from-pub-date:{min_year}-01-01",
    })
    items = data.get("message", {}).get("items", [])
    out: List[Ref] = []
    for it in items:
        title = (it.get("title") or [""])[0]
        year = str((it.get("issued", {}).get("date-parts", [[""]])[0] or [""])[0])
        doi = it.get("DOI", "")
        url = it.get("URL", "")
        authors_list = it.get("author", []) or []
        names = []
        for a in authors_list[:8]:
            last = a.get("family") or a.get("name") or ""
            first = a.get("given", "")
            names.append(f"{last}, {first}".strip(", "))
        authors = " and ".join(names)
        venue = (it.get("container-title") or [""])[0]
        base = re.sub(r"\W+", "", (authors_list[0].get("family") if authors_list else (title.split()[0] if title else "ref")).lower())
        yr = re.sub(r"\D", "", year)[:4] or "0000"
        first_word = re.sub(r"\W+", "", (title.split()[0] if title else "paper").lower())[:10]
        key = f"{base}{yr}_{first_word}"[:40]
        out.append(Ref(key=key, title=title, authors=authors, year=yr, venue=venue, doi=doi, url=url))
    return out

def collect_refs_for_chapter(chapter_title: str, key_points: List[str], per_topic: int = 3, budget: int = 12) -> List[Ref]:
    bag: List[Ref] = []
    seen = set()
    topics = [chapter_title] + list(key_points or [])
    for t in topics:
        refs = search_crossref(t, rows=per_topic)
        for r in refs:
            dup_key = r.doi or r.key
            if dup_key in seen:
                continue
            seen.add(dup_key)
            bag.append(r)
            if len(bag) >= budget:
                return bag
        time.sleep(0.3)  # polite pacing
    return bag

# =========================
# Orchestration
# =========================
class Checkpointer:
    def __init__(self, path: Path):
        self.path = path
        self.path.parent.mkdir(parents=True, exist_ok=True)
        if not self.path.exists():
            self.path.write_text("")
    def log(self, event: Dict[str, Any]):
        event["ts"] = datetime.now(timezone.utc).isoformat()
        with self.path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")

class BookBuilder:
    def __init__(self, llm: LLMClient, outdir: Path, idea: str, genre: str, audience: str,
                 reading_level: str, style: str, min_words: int, max_words: int, resume: bool = False,
                 temperature: float = 1.0, model: str = "", bibliography: Optional[Dict[str, Dict[str, Any]]] = None):
        self.llm = llm
        self.outdir = outdir
        self.plan_dir = outdir / "plan"
        self.chapters_dir = outdir / "chapters"
        self.idea = idea
        self.genre = genre
        self.audience = audience
        self.reading_level = reading_level
        self.style = style
        self.min_words = min_words
        self.max_words = max_words
        self.temperature = temperature
        self.model = model
        self.checkpoint = Checkpointer(outdir / "checkpoints.jsonl")
        self.plan_dir.mkdir(parents=True, exist_ok=True)
        self.chapters_dir.mkdir(parents=True, exist_ok=True)
        # State
        self.meta: Optional[BookMeta] = None
        self.plan: Optional[BookPlan] = None
        self.bibliography: Dict[str, Dict[str, Any]] = bibliography or {}
        self.auto_cite: bool = False
        if resume:
            self._try_load_state()

    def _try_load_state(self):
        meta_path = self.plan_dir / "meta.json"
        plan_path = self.plan_dir / "plan.json"
        if meta_path.exists():
            self.meta = BookMeta.model_validate_json(meta_path.read_text(encoding="utf-8"))
        if plan_path.exists():
            self.plan = BookPlan.model_validate_json(plan_path.read_text(encoding="utf-8"))

    def _save_json(self, path: Path, data: BaseModel | Dict[str, Any]):
        payload = data.model_dump() if isinstance(data, BaseModel) else data
        text = json.dumps(payload, indent=2, ensure_ascii=False)
        path.write_text(text, encoding="utf-8")

    # ----- Steps -----
    def generate_meta(self) -> BookMeta:
        if self.meta:
            return self.meta

        # Strict JSON first (OpenAI structured)
        meta_obj: dict = {}
        if isinstance(self.llm, OpenAIClient):
            try:
                meta_obj = self.llm.structured(
                    SYSTEM_POLICY,
                    META_PROMPT.format(
                        idea=self.idea, genre=self.genre, audience=self.audience,
                        reading_level=self.reading_level, style=self.style
                    ),
                    json_schema=META_SCHEMA,
                    max_tokens=800,
                    model=self.model,
                )
            except Exception as e:
                console.print(f"[yellow]structured meta failed, falling back: {e}")

        if not meta_obj:
            text = self.llm.complete(
                SYSTEM_POLICY,
                META_PROMPT.format(
                    idea=self.idea, genre=self.genre, audience=self.audience,
                    reading_level=self.reading_level, style=self.style
                ),
                max_tokens=1200,
                model=self.model,
            )
            (self.plan_dir / "meta_raw.txt").write_text(text, encoding="utf-8")
            try:
                meta_obj = _parse_or_repair_json(text)
            except Exception:
                meta_obj = {}

        def _f(x, d): return (x or "").strip() or d
        meta_obj = {
            "title": _f(meta_obj.get("title"), "Untitled Book"),
            "subtitle": _f(meta_obj.get("subtitle"), "A Project Generated by AI"),
            "author": _f(meta_obj.get("author"), "Anonymous"),
            "genre": _f(meta_obj.get("genre"), self.genre),
            "audience": _f(meta_obj.get("audience"), self.audience),
            "reading_level": _f(meta_obj.get("reading_level"), self.reading_level),
            "style": _f(meta_obj.get("style"), self.style),
            "promise": _f(meta_obj.get("promise"), "A fresh perspective created from a single idea."),
        }
        self.meta = BookMeta(**meta_obj)
        self._save_json(self.plan_dir / "meta.json", self.meta)
        self.checkpoint.log({"stage": "meta", "meta": self.meta.model_dump()})
        return self.meta

    def generate_plan(self) -> BookPlan:
        if self.plan:
            return self.plan
        meta = self.generate_meta()
        prompt = PLAN_PROMPT.format(
            meta_json=json.dumps(meta.model_dump(), indent=2, ensure_ascii=False),
            idea=self.idea,
            min_words=self.min_words,
            max_words=self.max_words,
        )

        plan_obj: dict = {}
        if isinstance(self.llm, OpenAIClient):
            try:
                plan_obj = self.llm.structured(
                    SYSTEM_POLICY,
                    prompt,
                    json_schema=PLAN_SCHEMA,
                    max_tokens=4000,
                    model=self.model,
                )
            except Exception as e:
                console.print(f"[yellow]structured plan failed, falling back: {e}")

        if not plan_obj:
            text = self.llm.complete(SYSTEM_POLICY, prompt, max_tokens=4000, model=self.model)
            (self.plan_dir / "plan_raw.txt").write_text(text, encoding="utf-8")
            try:
                plan_obj = _parse_or_repair_json(text)
            except Exception:
                strict = prompt + "\n\nReturn ONLY valid JSON. No prose outside the JSON."
                text2 = self.llm.complete(SYSTEM_POLICY, strict, max_tokens=5000, model=self.model)
                (self.plan_dir / "plan_raw_retry.txt").write_text(text2, encoding="utf-8")
                plan_obj = _parse_or_repair_json(text2)

        chs = plan_obj.get("chapters") or []
        if not isinstance(chs, list) or len(chs) < 8:
            raise RuntimeError("Plan invalid (too few chapters). See plan/plan_raw*.txt")

        self.plan = BookPlan(**plan_obj)
        self._save_json(self.plan_dir / "plan.json", self.plan)
        (self.plan_dir / "synopsis.md").write_text(self.plan.synopsis or "", encoding="utf-8")
        self.checkpoint.log({"stage": "plan", "total_target_words": self.plan.total_target_words})
        return self.plan

    def draft_chapter(self, ch: ChapterPlan, prev_titles: List[str]) -> str:
        meta = self.generate_meta()
        plan = self.generate_plan()

        # Auto-cite: fetch references per chapter and merge into bibliography
        if self.auto_cite:
            refs = collect_refs_for_chapter(ch.title, ch.key_points, per_topic=3, budget=12)
            # merge
            for r in refs:
                self.bibliography[r.key] = {
                    "title": r.title, "author": r.authors, "year": r.year,
                    "journal": r.venue, "doi": r.doi, "url": r.url,
                }
            # optional per-chapter .bib for inspection
            (self.plan_dir / f"auto_refs_ch{ch.number:02d}.bib").write_text(
                "".join(r.to_bibtex() for r in refs), encoding="utf-8"
            )

        bib_keys = "\n".join(
            f"- {v.get('title','(untitled)')} → {k}" for k, v in self.bibliography.items()
        ) or "(none provided)"

        CHAPTER_DRAFT_PROMPT = """
Draft Chapter {ch_num}: "{ch_title}" for the book described below. Write in the established voice.
Use the chapter plan faithfully, covering each section in order with smooth transitions.
Target ~{target_words} words (±10%).
Include:
- A short *italicized* opener (1–2 sentences).
- Clear section headings using Markdown H3 (### Section Title).
- Optional figure/table placeholders (e.g., *[Figure: ...]*), but do not invent data.
- End with 3–5 reflective questions or key takeaways as a bulleted list.

Citations (when writing a textbook):
- Cite only sources present in the bibliography list below, using the Markdown form `[@key]`.
- If no appropriate source exists in the bibliography, write without a citation rather than fabricating one.

Book metadata: {meta_json}
Book synopsis: {synopsis}
Chapter plan JSON: {chapter_json}

Available bibliography keys (title → key):
{bib_keys}

Previously drafted chapter titles:
{prev_titles}
""".strip()

        prompt = CHAPTER_DRAFT_PROMPT.format(
            ch_num=ch.number,
            ch_title=ch.title,
            target_words=ch.target_words,
            meta_json=json.dumps(meta.model_dump(), indent=2, ensure_ascii=False),
            synopsis=plan.synopsis,
            chapter_json=json.dumps(ch.model_dump(), indent=2, ensure_ascii=False),
            prev_titles="\n".join(f"- {t}" for t in prev_titles) if prev_titles else "(none)",
            bib_keys=bib_keys,
        )

        text = self.llm.complete(SYSTEM_POLICY, prompt, max_tokens=5000, model=self.model).strip()
        if len(text) < 300:
            # Retry once with a nudge; omit temp hassles
            prompt_retry = prompt + "\n\nWrite the full chapter now. Return prose, not JSON."
            text = self.llm.complete(SYSTEM_POLICY, prompt_retry, max_tokens=5500, model=self.model).strip()
            if len(text) < 300:
                raise RuntimeError(f"Chapter {ch.number} draft came back empty twice.")
        return text

    def run(self):
        meta = self.generate_meta()
        plan = self.generate_plan()

        chapter_files: List[Path] = []
        prev_titles: List[str] = []

        for ch in tqdm(plan.chapters, desc="Drafting chapters"):
            ch_name = f"{ch.number:02d}_{self._sanitize_filename(ch.title)}.md"
            ch_path = self.chapters_dir / ch_name
            if ch_path.exists() and ch_path.stat().st_size > 200:
                prev_titles.append(ch.title)
                chapter_files.append(ch_path)
                continue

            text = self.draft_chapter(ch, prev_titles)
            ch_doc = f"# {meta.title}: {ch.title}\n\n{text.strip()}\n"
            ch_path.write_text(ch_doc, encoding="utf-8")
            self.checkpoint.log({"stage": "chapter", "chapter": ch.number, "title": ch.title, "path": str(ch_path)})
            prev_titles.append(ch.title)
            chapter_files.append(ch_path)

        if not chapter_files:
            raise RuntimeError("No chapters were drafted. Check plan/plan_raw*.txt for JSON issues.")

        book_md = self._assemble_book_md(meta, plan, chapter_files)
        (self.outdir / "book.md").write_text(book_md, encoding="utf-8")
        self.checkpoint.log({"stage": "assemble", "path": str(self.outdir / "book.md")})

        # Optional exports
        try:
            self._export_epub(meta, book_md)
        except Exception as e:
            console.print(f"[yellow]EPUB export skipped/failed: {e}")
        try:
            self._export_docx(book_md)
        except Exception as e:
            console.print(f"[yellow]DOCX export skipped/failed: {e}")
        try:
            self._export_pdf(book_md)
        except Exception as e:
            console.print(f"[yellow]PDF export skipped/failed: {e}")

        console.print("\n[bold green]Done![/] Manuscript ready at:", self.outdir / "book.md")

    # ----- Helpers -----
    @staticmethod
    def _sanitize_filename(name: str) -> str:
        s = re.sub(r"[^A-Za-z0-9\- _]+", "", name).strip().replace(" ", "_")
        return s[:120] if s else "chapter"

    def _assemble_book_md(self, meta: BookMeta, plan: BookPlan, chapter_files: List[Path]) -> str:
        front = [
            f"# {meta.title}",
            f"## {meta.subtitle}",
            f"**Author:** {meta.author}",
            "\n---\n",
            "## About this book",
            (plan.synopsis or "").strip(),
            "\n---\n",
            "## Table of Contents",
        ]
        for ch in plan.chapters:
            front.append(f"- Chapter {ch.number}: {ch.title}")
        front.append("\n---\n")

        parts = [p.read_text(encoding="utf-8") for p in chapter_files]
        manuscript = "\n\n".join(front + parts) + "\n"

        # Append References chapter (only cited keys)
        if self.bibliography:
            cited = self._extract_citekeys(manuscript)
            if cited:
                bib_md = self._format_bibliography(cited)
                manuscript += "\n\n---\n\n# References\n\n" + bib_md + "\n"
        return manuscript

    def _extract_citekeys(self, text: str) -> List[str]:
        keys = set(re.findall(r"\[@([^\]]+)\]", text))
        expanded = set()
        for k in keys:
            for part in re.split(r"[;,]\s*", k):
                if part:
                    expanded.add(part.strip(" @"))
        return sorted(expanded)

    def _format_bibliography(self, keys: List[str]) -> str:
        lines = []
        for k in keys:
            entry = self.bibliography.get(k)
            if not entry:
                lines.append(f"- **MISSING** citation for key `{k}` — please supply in bibliography.")
                continue
            authors = entry.get("author") or entry.get("authors") or ""
            year = entry.get("year") or entry.get("date", "")[0:4]
            title = entry.get("title", "Untitled")
            journal = entry.get("journal") or entry.get("booktitle") or entry.get("publisher") or ""
            doi = entry.get("doi", "")
            url = entry.get("url", "")
            tail = f" {journal}" if journal else ""
            tail += f". DOI: {doi}" if doi else ""
            tail += f". {url}" if url else ""
            lines.append(f"- {authors} ({year}). *{title}*.{tail}")
        return "\n".join(lines)

    # ----- Exports -----
    def _export_epub(self, meta: BookMeta, book_md: str):
        if epub is None:
            return
        bk = epub.EpubBook()
        bk.set_identifier("ai-ebook-" + re.sub(r"\W+", "-", meta.title.lower()))
        bk.set_title(meta.title)
        bk.set_language("en")
        bk.add_author(meta.author)
        chunks = re.split(r"\n(?=# )", book_md)
        spine = []
        toc = []
        for i, chunk in enumerate(chunks, 1):
            html = markdown.markdown(chunk) if markdown else f"<pre>{chunk}</pre>"
            c = epub.EpubHtml(title=f"Section {i}", file_name=f"sec_{i}.xhtml", content=html)
            bk.add_item(c)
            spine.append(c)
            toc.append(c)
        bk.toc = tuple(toc)
        bk.spine = ["nav"] + spine
        bk.add_item(epub.EpubNcx())
        bk.add_item(epub.EpubNav())
        out = self.outdir / "book.epub"
        epub.write_epub(str(out), bk)
        console.print(f"[green]EPUB written:[/] {out}")

    def _export_docx(self, book_md: str):
        if Document is None:
            return
        from markdown import markdown as md_to_html
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except Exception:
            console.print("[yellow]DOCX export needs beautifulsoup4; skipping.")
            return
        html = md_to_html(book_md)
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        for node in soup.children:
            txt = node.get_text(strip=False)
            if node.name == "h1":
                doc.add_heading(txt, level=0)
            elif node.name == "h2":
                doc.add_heading(txt, level=1)
            elif node.name == "h3":
                doc.add_heading(txt, level=2)
            elif node.name == "p":
                doc.add_paragraph(txt)
            elif node.name == "ul":
                for li in node.find_all("li", recursive=False):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            else:
                if txt.strip():
                    doc.add_paragraph(txt)
        out = self.outdir / "book.docx"
        doc.save(str(out))
        console.print(f"[green]DOCX written:[/] {out}")

    def _export_pdf(self, book_md: str):
        try:
            from reportlab.lib.pagesizes import LETTER
            from reportlab.pdfgen import canvas as pdf_canvas
            from reportlab.lib.units import inch
        except Exception:
            # Optional: use Pandoc/LaTeX for higher quality PDFs
            return
        out_path = self.outdir / "book.pdf"
        c = pdf_canvas.Canvas(str(out_path), pagesize=LETTER)
        width, height = LETTER
        margin = 0.75 * inch
        y = height - margin
        for line in book_md.splitlines():
            if not line.strip():
                y -= 14
                if y < margin:
                    c.showPage(); y = height - margin
                continue
            words = line.split(" ")
            cur = ""
            for w in words:
                test = (cur + " " + w).strip()
                if len(test) > 95:
                    c.drawString(margin, y, cur)
                    y -= 12
                    if y < margin:
                        c.showPage(); y = height - margin
                    cur = w
                else:
                    cur = test
            if cur:
                c.drawString(margin, y, cur)
                y -= 12
                if y < margin:
                    c.showPage(); y = height - margin
        c.showPage(); c.save()
        console.print(f"[green]PDF written:[/] {out_path}")

# =========================
# CLI
# =========================
def build_llm(backend: str, model: str) -> LLMClient:
    backend = backend.lower()
    if backend == "openai":
        return OpenAIClient(model=model)
    if backend == "anthropic":
        return AnthropicClient(model=model)
    if backend in {"ollama", "local"}:
        return OllamaClient(model=model)
    raise ValueError(f"Unsupported backend: {backend}")

def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate a full e-book from a single idea using an LLM.")
    p.add_argument("--idea", type=str, required=False, default="An untitled book idea",
                   help="Single-sentence core idea for the book.")
    p.add_argument("--genre", type=str, default="non-fiction")
    p.add_argument("--audience", type=str, default="general audience")
    p.add_argument("--reading-level", type=str, default="grade 10")
    p.add_argument("--style", type=str, default="clear, engaging, practical")
    p.add_argument("--min-words", type=int, default=105_000)
    p.add_argument("--max-words", type=int, default=125_000)
    p.add_argument("--backend", type=str, choices=["openai", "anthropic", "ollama", "local"], default="openai")
    p.add_argument("--model", type=str, default="gpt-4.1")
    p.add_argument("--temperature", type=float, default=1.0)
    p.add_argument("--outdir", type=Path, default=Path("./book_out"))
    p.add_argument("--resume", action="store_true", help="Resume from existing plan/chapters if present.")
    p.add_argument("--no-exports", action="store_true", help="Skip EPUB/DOCX/PDF exports; keep Markdown only.")
    # Bibliography / auto-cite
    p.add_argument("--bibtex", type=Path, default=None, help="Optional path to a .bib file of sources.")
    p.add_argument("--citations-json", type=Path, default=None, help="Optional path to citations JSON (list or {key: entry}).")
    p.add_argument("--auto-cite", action="store_true", help="Auto-search literature per chapter (Crossref).")
    return p.parse_args(argv)

def main(argv: Optional[List[str]] = None):
    args = parse_args(argv)
    args.outdir.mkdir(parents=True, exist_ok=True)

    llm = build_llm(args.backend, args.model)

    # Load user-provided bibliography (optional)
    bibliography: Dict[str, Dict[str, Any]] = {}
    if args.bibtex and args.bibtex.exists():
        if bibtexparser is None:
            console.print("[yellow]BibTeX provided but bibtexparser not installed. Skipping.")
        else:
            with open(args.bibtex, "r", encoding="utf-8") as bf:
                db = bibtexparser.load(bf)
            for entry in db.entries:
                key = entry.get("ID") or entry.get("key") or entry.get("citation_key")
                if key:
                    bibliography[key] = entry
    if args.citations_json and args.citations_json.exists():
        try:
            cj = json.loads(args.citations_json.read_text(encoding="utf-8"))
            if isinstance(cj, list):
                for i, e in enumerate(cj, 1):
                    key = e.get("key") or e.get("id") or f"ref{i:04d}"
                    bibliography[key] = e
            elif isinstance(cj, dict):
                bibliography.update({str(k): v for k, v in cj.items()})
        except Exception as e:
            console.print(f"[yellow]Failed to read citations JSON: {e}")

    builder = BookBuilder(
        llm=llm,
        outdir=args.outdir,
        idea=args.idea,
        genre=args.genre,
        audience=args.audience,
        reading_level=args.reading_level,
        style=args.style,
        min_words=args.min_words,
        max_words=args.max_words,
        resume=args.resume,
        temperature=args.temperature,
        model=args.model,
        bibliography=bibliography,
    )
    builder.auto_cite = args.auto_cite

    builder.run()

if __name__ == "__main__":
    main()
