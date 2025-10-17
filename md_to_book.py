#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_book_legacy_like.py — ToC-singleton edition (v2)
Fix: aggressively removes any pre-existing ToC, including
inline "Table of Contents — Chapter 1 — Chapter 2 …" paragraphs.
"""
from __future__ import annotations
import argparse, re
from pathlib import Path

try:
    import markdown
except Exception:
    markdown = None

try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

try:
    from ebooklib import epub
except Exception:
    epub = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Preformatted
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
except Exception:
    SimpleDocTemplate = None


def read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8")


def write_text(p: Path, s: str):
    p.write_text(s, encoding="utf-8")


def clean_markdown(md: str, strip_citations: bool, strip_figures: bool) -> str:
    # Fix broken hyphenation across line breaks
    md = re.sub(r"([A-Za-z])-\n([A-Za-z])", r"\1\2", md)
    # Replace spaced hyphen with em dash
    md = re.sub(r"\s+-\s+(?=\w)", " — ", md)
    if strip_citations:
        md = re.sub(r"\[@[^\]]+\]", "", md)  # remove [@keys]
    if strip_figures:
        md = re.sub(r'^\s*!\[[^\]]*\]\([^)]+\)\s*$', '', md, flags=re.M)
        md = re.sub(r'\*\[Figure:[^\]]+\]\*', '', md)
    md = re.sub(r"\n{3,}", "\n\n", md).strip() + "\n"
    return md


def parse_h1_list(md: str):
    heads = re.findall(r'(?m)^#\s+(.+)$', md)
    # drop the very first H1 (title), keep subsequent H1s (chapters)
    return heads[1:] if len(heads) > 1 else []


def build_toc(md: str) -> str:
    items = []
    for i, full in enumerate(parse_h1_list(md), 1):
        title = full.split(":", 1)[1].strip() if ":" in full else full.strip()
        items.append(f"- Chapter {i}: {title}")
    return "\n".join(items) if items else "- (no chapters found)"


def strip_all_toc_blocks(md: str) -> str:
    """
    Remove any existing ToC blocks, including:
      1) '## Table of Contents' section (normal case)
      2) '# Table of Contents' section (rare)
      3) Inline paragraph starting with 'Table of Contents — Chapter ...'
    """
    # 1) & 2) section-form ToC (H2/H1)
    sec_pat = re.compile(
        r'(?ms)^\s{0,3}(?:##|#)\s+Table of Contents\s*$.*?(?=^\s{0,3}(?:##|#)\s+|^\s*---\s*$|\Z)'
    )
    md = re.sub(sec_pat, '', md)

    # 3) inline, single-paragraph ToC (em dashes with Chapter labels)
    inline_pat = re.compile(
        r'(?ms)^\s{0,3}Table of Contents\b[^\n]*?—\s*Chapter\s+1\b.*?(?=^\s{0,3}(?:##|#)\s+|^\s*---\s*$|\Z)'
    )
    md = re.sub(inline_pat, '', md)

    # clean big gaps
    md = re.sub(r'\n{3,}', '\n\n', md).strip() + "\n"
    return md


def insert_single_toc(md: str) -> str:
    bullets = build_toc(md)
    block = f"\n\n## Table of Contents\n{bullets}\n\n"

    about = re.search(r'(?m)^##\s+About this book\s*$', md)
    if about:
        start = about.end()
        tail = md[start:]
        m_end = re.search(r'(?m)^(?:---|##\s+|#\s+)', tail)
        ins = start + (m_end.start() if m_end else 0)
        return md[:ins] + block + md[ins:]

    h1 = re.search(r'(?m)^#\s+.+$', md)
    if h1:
        return md[:h1.end()] + block + md[h1.end():]
    return block + md


def rebuild_singleton_toc(md: str, enable: bool) -> str:
    if not enable:
        return md
    md = strip_all_toc_blocks(md)
    return insert_single_toc(md)


def split_by_h1(md: str):
    parts = re.split(r'(?m)^(# .+)\n', md)
    out = []
    if not parts or len(parts) == 1:
        return [("Manuscript", md)]
    preface = parts[0].strip()
    i = 1
    if preface:
        out.append(("Front Matter", preface))
    while i < len(parts):
        h = parts[i].strip()
        body = parts[i+1] if i+1 < len(parts) else ""
        title = h.lstrip("# ").strip()
        out.append((title, body))
        i += 2
    return out


def to_html(md: str) -> str:
    if markdown is None:
        safe = md.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return f"<pre>{safe}</pre>"
    return markdown.markdown(md, extensions=["fenced_code", "tables"]).strip()


def make_epub(md: str, outdir: Path, title: str, author: str) -> Path:
    if epub is None:
        return outdir / "book.epub"
    bk = epub.EpubBook()
    import re as _re
    bk.set_identifier("md-book-" + _re.sub(r"\W+", "-", (title or 'book').lower()))
    bk.set_title(title or "Book")
    bk.set_language("en")
    if author:
        bk.add_author(author)

    chapters = []
    for idx, (sec_title, chunk) in enumerate(split_by_h1(md), 1):
        html = to_html(f"# {sec_title}\n\n{chunk}")
        item = epub.EpubHtml(title=sec_title, file_name=f"sec_{idx:02d}.xhtml", content=html)
        bk.add_item(item)
        chapters.append(item)

    bk.toc = tuple(chapters)
    bk.spine = ["nav"] + chapters
    bk.add_item(epub.EpubNcx())
    bk.add_item(epub.EpubNav())

    out = outdir / "book.epub"
    epub.write_epub(str(out), bk)
    return out


def make_docx(md: str, outdir: Path, title: str, author: str) -> Path:
    if Document is None or BeautifulSoup is None or markdown is None:
        return outdir / "book.docx"
    html = to_html(md)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    if author:
        doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph("")
    for node in soup.children:
        if getattr(node, "name", None) is None:
            continue
        txt = node.get_text(strip=False)
        if node.name == "h1":
            doc.add_heading(txt, level=1)
        elif node.name == "h2":
            doc.add_heading(txt, level=2)
        elif node.name == "h3":
            doc.add_heading(txt, level=3)
        elif node.name == "p":
            doc.add_paragraph(txt)
        elif node.name == "ul":
            for li in node.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif node.name == "ol":
            for li in node.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")
        else:
            if txt.strip():
                doc.add_paragraph(txt)
    out = outdir / "book.docx"
    doc.save(str(out))
    return out


def make_pdf(md: str, outdir: Path, title: str, author: str) -> Path:
    out = outdir / "book.pdf"
    if SimpleDocTemplate is None:
        return out
    doc = SimpleDocTemplate(str(out), pagesize=LETTER,
                            leftMargin=0.85*inch, rightMargin=0.85*inch,
                            topMargin=0.85*inch, bottomMargin=0.85*inch)
    styles = getSampleStyleSheet()
    H1 = ParagraphStyle('H1', parent=styles['Heading1'], alignment=TA_LEFT, spaceAfter=8)
    H2 = ParagraphStyle('H2', parent=styles['Heading2'], alignment=TA_LEFT, spaceAfter=6)
    H3 = ParagraphStyle('H3', parent=styles['Heading3'], alignment=TA_LEFT, spaceAfter=4)
    P  = ParagraphStyle('P',  parent=styles['BodyText'],  alignment=TA_LEFT, leading=14, spaceAfter=6)
    I  = ParagraphStyle('I',  parent=P, italic=True)

    story = []
    if title:
        story.append(Paragraph(title, H1))
    if author:
        story.append(Paragraph(f"Author: {author}", P))
    story.append(Spacer(1, 12))

    html = to_html(md)
    if BeautifulSoup is None:
        story.append(Preformatted(md, P))
        doc.build(story)
        return out

    soup = BeautifulSoup(html, "html.parser")
    for node in soup.children:
        name = getattr(node, "name", None)
        if not name:
            continue
        txt = node.get_text(strip=False).replace("\n", "<br/>")
        if name == "h1":
            story.append(Paragraph(txt, H1))
        elif name == "h2":
            story.append(Paragraph(txt, H2))
        elif name == "h3":
            story.append(Paragraph(txt, H3))
        elif name == "p":
            story.append(Paragraph(txt, I if node.find('em') else P))
        elif name in ("ul", "ol"):
            for li in node.find_all("li", recursive=False):
                story.append(Paragraph("• " + li.get_text(), P))
        else:
            story.append(Paragraph(txt, P))
        story.append(Spacer(1, 6))
    doc.build(story)
    return out


def main():
    ap = argparse.ArgumentParser(description="Convert Markdown to EPUB/DOCX/PDF and keep a single, rebuilt ToC.")
    ap.add_argument("--input", type=Path, required=True)
    ap.add_argument("--outdir", type=Path, required=True)
    ap.add_argument("--title", type=str, default="")
    ap.add_argument("--author", type=str, default="")
    ap.add_argument("--strip-citations", action="store_true")
    ap.add_argument("--strip-figures", action="store_true")
    ap.add_argument("--no-rebuild-toc", action="store_true", help="Do not rebuild ToC (keeps whatever is in the MD).")
    ap.add_argument("--all", action="store_true")
    args = ap.parse_args()

    args.outdir.mkdir(parents=True, exist_ok=True)
    md = read_text(args.input)
    md = clean_markdown(md, args.strip_citations, args.strip_figures)

    if not args.no_rebuild_toc:
        md = rebuild_singleton_toc(md, enable=True)

    cleaned = args.outdir / "book.cleaned.md"
    write_text(cleaned, md)

    epub_p = make_epub(md, args.outdir, args.title or "Book", args.author or "Anonymous")
    docx_p = make_docx(md, args.outdir, args.title, args.author)
    pdf_p  = make_pdf(md, args.outdir, args.title, args.author)

    print(f"Wrote: {epub_p}")
    print(f"Wrote: {docx_p}")
    print(f"Wrote: {pdf_p}")


if __name__ == "__main__":
    main()
